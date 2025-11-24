# -*- coding: utf-8 -*-
import json
import re
import os
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def safe_image(minimal_data, key, width=400, height=200):
    path = minimal_data.get("image_paths", {}).get(key)
    if not path:
        # Если изображения нет, возвращаем плейсхолдер
        path = "placeholder.png"  # <-- файл-заглушка, например белый фон или "нет изображения"
    return {"path": path, "width": width, "height": height}


# JSON data with <b>...</b> tags for bold formatting
def build_report_data(minimal_data):
    """
    Builds the full report_data structure from minimal input JSON.
    Minimal JSON example:
    {
        "period": {
            "start_date": "12.04.2025",
            "end_date": "25.04.2025"
        },
        "udsh_measurements": [
            {"party": 1, "registered": 2628},
            {"party": 2, "registered": 3324},
            {"party": 3, "registered": 3847}
        ],
        "image_paths": {
            "image1": "image1.png",
            "image2": "image2.png",
            "image3": "image3.png",
            "image4": "image4.jpeg",
            "image5": "image5.jpeg",
            "image6": "image6.png"
        }
    }
    """
    # Default values for other fields
    default_data = {
        "report_type": "КЗ201",
        "title": {
            "text": "ОТЧЕТ<b>КЗ201 </b>\n\nОБ <b>ЭКСПЛУАТАЦИИ КРИОГЕННОГО ЗАМЕДЛИТЕЛЯ</b>",
            "font_size": 14,
            "alignment": "center"
        },
        "subtitle": {
            "text": "",
            "font_size": 12,
            "alignment": "center"
        },
        "period": minimal_data["period"],
        "sections": [
            {
                "id": 1,
                "title": {"text": "Цикл:", "font_size": 10},
                "content": {"text": f"<b>{minimal_data['period']['start_date']} – {minimal_data['period']['end_date']}</b>", "font_size": 10}
            },
            {
                "id": 2,
                "title": {"text": "Состояние реактора:", "font_size": 10},
                "content": {"text": "на мощности <b>1,4 кВт</b>", "font_size": 10}
            },
            {
                "id": 3,
                "title": {"text": "Технологическая схема:", "font_size": 10},
                "content": {"text": "В приложении 1", "font_size": 10}
            },
            {
                "id": 4,
                "title": {"text": "Загрузка шариков", "font_size": 10},
                "subsections": [
                    {
                        "title": {"text": "Фактически загружено:", "font_size": 10},
                        "table": {
                            "headers": ["Партия<", "Объем, мл"],
                            "rows": [["1", "<b>330</b>"], ["2", "<b>330</b>"], ["3", "<b>330</b>"]],
                            "font_size": 10,
                            "bold_columns": [],
                            "align": "left"
                        }
                    },
                    {
                        "title": {"text": "Измерено при помощи УДШ (Узел детектирования шариков):", "font_size": 10},
                        "table": {
                            "headers": ["Партия", "Зарегистрировано, шт.", "Теоретическое кол-во, шт", "Погрешность фактическая, %"],
                            "rows": [],
                            "font_size": 10,
                            "bold_columns": [],
                            "align": "left"
                        }
                    }
                ]
            },
            {
                "id": 5,
                "title": {"text": "Показания газгольдера", "font_size": 10},
                "content": [
                    {
                        "text": {"text": "Рисунок 1. Показания с <b>ЛИР</b>", "font_size": 9},
                        "image": safe_image(minimal_data, "image1")
                    }
                ]
            },
            {
                "id": 6,
                "title": {"text": "Показания температур", "font_size": 10},
                "content": [
                    {
                        "text": {"text": "Рисунок 2. Температурные данные с термодиода <b>DT_51</b>", "font_size": 9},
                        "image": safe_image(minimal_data, "image2")
                    },
                    {
                        "text": {"text": "Рисунок 4. Температурные данные с термопары <b>Т32</b> в период с 9.00 до 10.00 14.04.2025г.", "font_size": 9},
                        "image": safe_image(minimal_data, "image4")
                    }
                ]
            },
            {
                "id": 7,
                "title": {"text": "Показания давления (вакуума)", "font_size": 10},
                "content": [
                    {
                        "text": {"text": "Рисунок 3. Показания вакуумметрического датчика <b>P22</b> в контуре", "font_size": 9},
                        "image": safe_image(minimal_data, "image3")
                    },
                    {
                        "text": {"text": "Рисунок 5. Показания вакуумметрического датчика <b>ВД21</b> в рубашке.", "font_size": 9},
                        "image": safe_image(minimal_data, "image5")
                    }
                ]
            },
            {
                "id": 8,
                "title": {"text": "Контроль концентрации радиолитического водорода и кислорода при отогреве", "font_size": 10},
                "subsections": [
                    {
                        "title": {"text": "Контроль концентрации радиолитического водорода", "font_size": 10},
                        "content": {"text": "<b>Не проводилось</b>", "font_size": 10}
                    },
                    {
                        "title": {"text": "Контроль концентрации кислорода", "font_size": 10},
                        "content": {"text": "<b>Не проводилось</b>", "font_size": 10}
                    }
                ]
            },
            {
                "id": 9,
                "title": {"text": "Мезитилен", "font_size": 10},
                "subsections": [
                    {
                        "title": {"text": "Объем слитого мезитилена:", "font_size": 10},
                        "table": {
                            "headers": ["Дата", "<b>Объем, мл</b>", "Примечания"],
                            "rows": [
                                ["05.05.2025", "<b>610</b>", "Из камеры"],
                                ["12.05.2025", "<b>10</b>", "Из внутреннего трубопровода"],
                                ["15.05.2025", "<b>5</b>", "Из внутреннего трубопровода"],
                                ["Всего:", "<b>625</b>", "-"]
                            ],
                            "font_size": 10,
                            "bold_columns": [],
                            "align": "left"
                        }
                    },
                    {
                        "title": {"text": "Измерение вязкости мезитилена:", "font_size": 10},
                        "content": {"text": "Вязкость слитого мезитилена <b>КЗ 201</b> - <b>23 сР</b>", "font_size": 10}
                    }
                ]
            },
            {
                "id": 10,
                "title": {"text": "Дополнительные измерения", "font_size": 10},
                "content": [
                    {"text": "Излучение по <b>гамма</b>: <b>19 мкЗв</b>", "font_size": 10},
                    {"text": "Излучение по <b>бетта</b>: <b>2200</b> частиц", "font_size": 10}
                ]
            },
            {
                "id": 11,
                "title": {"text": "Результаты", "font_size": 10},
                "content": minimal_data["content"]
            },
            {
                "id": 12,
                "title": {"text": "Выводы", "font_size": 10},
                "content": {"text": "", "font_size": 10}
            },
            {
                "id": 13,
                "title": {"text": "Заключение (рекомендации)", "font_size": 10},
                "content": {"text": "", "font_size": 10}
            }
        ],
        "developers": [
            {"role": "Начальник группы №2 Сектор НИиКЗ ЛНФ ОИЯИ", "name": "Галушко <b>А.В.</b>", "signature": "       ", "font_size": 10}
        ],
        "approvers": [
            {"role": "Начальник сектора НИиКЗ ЛНФ ОИЯИ", "name": "Булавин <b>М.В.</b>", "signature": "       ", "font_size": 10},
            {"role": "Начальник МТО ЛНФ ОИЯИ", "name": "Слотвицкий <b>Ю.М.</b>", "signature": "       ", "font_size": 10},
            {"role": "Начальник группы №2 МТО ЛНФ ОИЯИ", "name": "Скуратов <b>В.А.</b>", "signature": "       ", "font_size": 10}
        ],
        "appendices": [
            {
                "title": {"text": "ПРИЛОЖЕНИЕ 1. <b>ТЕХНОЛОГИЧЕСКАЯ СХЕМА</b>", "font_size": 14},
                "content": {"image": {"path": minimal_data["image_paths"]["image6"], "width": 600, "height": 400}}
            },
            {
                "title": {"text": "ПРИЛОЖЕНИЕ 2. <b>ПРОГРАММА РАБОТ</b>", "font_size": 14},
                "content": {"text": "", "font_size": 10}
            }
        ]
    }

    # Integrate UDSH measurements into the table
    udsh_table = default_data["sections"][3]["subsections"][1]["table"]
    theoretical = 9000

    for meas in minimal_data["udsh_measurements"]:
        party = f"{meas['party']}"
        registered = f"<b>{meas['registered']}</b>"  # ← жирное выделение через HTML
        error = f"{((theoretical - meas['registered']) / theoretical * 100):.1f}"
        udsh_table["rows"].append([party, registered, str(theoretical), error])

    return default_data


# Sample minimal JSON for testing (as requested)
sample_minimal_json = {
    "period": {
        "start_date": "01.10.2025",
        "end_date": "13.10.2025"
    },
    "udsh_measurements": [
        {"party": 1, "registered": 2800},
        {"party": 2, "registered": 3500},
        {"party": 3, "registered": 4100}
    ],
    "image_paths": {
        "image1": "image1.png",
        "image2": "image2.png",
        "image3": "image3.png",
        "image4": "image4.jpeg",
        "image5": "image5.jpeg",
        "image6": r"C:\Users\Иван Литвак\source\repos\Автоматизация отчетов\Автоматизация отчетов\Техносхема.jpg"
    },

    "content": [
        {"text": "Все работы на <b>КЗ</b> выполнялись в соответствие с утвержденным Планом работ по криогенным замедлителям <b>КЗ201, КЗ202</b> и стенду КЗ201 на период <b>01.01.2025 – 30.06.2025г.</b> (см. приложение 2).", "font_size": 10},
        {"text": f"<b>Датаг.</b> выполнена подготовка КЗ201 к работе по инструкции, все системы работали штатно.", "font_size": 10},
        {"text": f"<b>Датаг.</b> - охлаждение КЗ прошло штатно.", "font_size": 10},
        {"text": f"<b>08:40-12.40 13.04.2025 г.</b> выполнена загрузка трёх партий шариков, загрузка прошла штатно. Общее время загрузки шариков в камеру составило 4 часов.", "font_size": 10},
        {"text": "Около <b>09:10 14.04.2025 г.</b> произошло отключение газодувки в связи с кратковременным обрывом связи <b>modbus</b> с газодувкой, внутри контура с гелием зафиксирован скачок температур до <b>136К</b> (см. рис.4) и падения вакуума в рубашке до <b>0,036 Торр</b> (см. рис.5).", "font_size": 10},
        {"text": "Остановка газодувки не сопровождалась ни выводом ошибки, ни звуковыми, ни световыми сигналами. Около <b>09:37 14.04.2025 г.</b> газодувка была запущена снова и все данные вернулись к штатным значениям.", "font_size": 10},
        {"text": "При отогреве произошел выброс <b>радиоактивности</b>…..", "font_size": 10}
    ]
}

import platform

if platform.system() == "Windows":
    font_path = r"C:\Windows\Fonts\times.ttf"
    font_bold_path = r"C:\Windows\Fonts\timesbd.ttf"
elif platform.system() == "Linux":
    font_path = "/usr/share/fonts/truetype/msttcorefonts/times.ttf"
    font_bold_path = "/usr/share/fonts/truetype/msttcorefonts/timesbd.ttf"
else:
    raise RuntimeError("Неизвестная ОС для регистрации шрифтов")

if not os.path.exists(font_path) or not os.path.exists(font_bold_path):
    raise FileNotFoundError("Шрифты Times New Roman не найдены")

pdfmetrics.registerFont(TTFont("Times New Roman", font_path))
pdfmetrics.registerFont(TTFont("Times New Roman Bold", font_bold_path))

pdfmetrics.registerFontFamily("Times New Roman", normal="Times New Roman", bold="Times New Roman Bold")




def create_pdf_report(data, output_path):
    doc = SimpleDocTemplate(output_path, pagesize=A4)
    styles = getSampleStyleSheet()

    # Dynamic style creation
    def get_paragraph_style(font_size, alignment="left"):
        name = f"CustomStyle_{font_size}_{alignment}"
        if name not in styles:
            align_map = {"left": 0, "center": 1, "right": 2}
            styles.add(ParagraphStyle(
                name=name,
                fontName="Times New Roman",
                fontSize=font_size,
                alignment=align_map.get(alignment, 0),
                spaceAfter=12,
                leading=font_size + 2,
                encoding='utf-8'
            ))
        return styles[name]

    def format_text(text):
        # text = str(text).replace('<b>', '<b><font name="TimesNewRoman-Bold">').replace('</b>', '</font></b>')
        return str(text)

    def format_table_cell(text, font_size, align="left"):
        text = format_text(text)
        return Paragraph(text, get_paragraph_style(font_size, align))

    def add_centered_image(elements, img_path, max_width=450):
        if os.path.exists(img_path):
            try:
                img_reader = ImageReader(img_path)
                orig_w, orig_h = img_reader.getSize()
                if orig_w > max_width:
                    scale_h = max_width * orig_h / orig_w
                    img = Image(img_path, width=max_width, height=scale_h)
                else:
                    img = Image(img_path, width=orig_w, height=orig_h)
                img_table = Table([[img]], colWidths=[img.width])
                img_table.setStyle(TableStyle([('ALIGN', (0,0),(-1,-1),'CENTER')]))
                elements.append(img_table)
            except Exception:
                elements.append(Paragraph(f"[Ошибка изображения: {img_path}]", get_paragraph_style(10)))
        else:
            elements.append(Paragraph(f"[Изображение не найдено: {img_path}]", get_paragraph_style(10)))

    elements = []

    # Title and subtitle
    for item in [data['title'], data['subtitle']]:
        text = format_text(item['text'])
        elements.append(Paragraph(text, get_paragraph_style(item['font_size'], item['alignment'])))
    elements.append(Spacer(1, 12))

    # Sections
    for section in data['sections']:
        title_text = format_text(f"{section['id']}. <b>{section['title']['text']}</b>")
        elements.append(Paragraph(title_text, get_paragraph_style(section['title']['font_size'])))
        elements.append(Spacer(1, 6))

        if 'content' in section:
            if isinstance(section['content'], dict) and 'text' in section['content']:
                text = format_text(section['content']['text'])
                elements.append(Paragraph(text, get_paragraph_style(section['content']['font_size'])))
            elif isinstance(section['content'], list):
                for item in section['content']:
                    if isinstance(item, dict) and 'image' in item:
                        add_centered_image(elements, item['image']['path'])
                        text = format_text(item['text']['text'])
                        elements.append(Paragraph(text, get_paragraph_style(item['text']['font_size'], "center")))
                    else:
                        text = format_text(item['text'])
                        elements.append(Paragraph(text, get_paragraph_style(item['font_size'])))

        if 'table' in section:
            table_data = [[format_table_cell(h, section['table']['font_size'], section['table']['align']) for h in section['table']['headers']]]
            for row in section['table']['rows']:
                table_data.append([format_table_cell(cell, section['table']['font_size'], section['table']['align']) for cell in row])
            table = Table(table_data)
            table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                ('FONTNAME', (0, 0), (-1, -1), 'Times New Roman'),
                ('FONTSIZE', (0, 0), (-1, -1), section['table']['font_size']),
                ('ALIGN', (0, 0), (-1, -1), section['table']['align'].upper()),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black)
            ]))
            elements.append(table)
            elements.append(Spacer(1, 6))

        if 'subsections' in section:
            for subsection in section['subsections']:
                title_text = format_text(subsection['title']['text'])
                elements.append(Paragraph(title_text, get_paragraph_style(subsection['title']['font_size'])))
                elements.append(Spacer(1, 6))
                if 'table' in subsection:
                    table_data = [[format_table_cell(h, subsection['table']['font_size'], subsection['table']['align']) for h in subsection['table']['headers']]]
                    for row in subsection['table']['rows']:
                        table_data.append([format_table_cell(cell, subsection['table']['font_size'], subsection['table']['align']) for cell in row])
                    table = Table(table_data)
                    table.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('FONTNAME', (0, 0), (-1, -1), 'Times New Roman'),
                        ('FONTSIZE', (0, 0), (-1, -1), subsection['table']['font_size']),
                        ('ALIGN', (0, 0), (-1, -1), subsection['table']['align'].upper()),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black)
                    ]))
                    elements.append(table)
                    elements.append(Spacer(1, 6))
                elif 'content' in subsection:
                    text = format_text(subsection['content']['text'])
                    elements.append(Paragraph(text, get_paragraph_style(subsection['content']['font_size'])))
                elements.append(Spacer(1, 6))
        elements.append(Spacer(1, 12))

    # Developers and Approvers Table
    for title_key, people in [("РАЗРАБОТАНО:", data['developers']), ("СОГЛАСОВАНО:", data['approvers'])]:
        elements.append(Paragraph(title_key, get_paragraph_style(10)))
        elements.append(Spacer(1, 6))
        table_data = [["<b>Должность</b>", "<b>Подпись</b>", "<b>Ф.И.О.</b>"]]
        for p in people:
            row = [p['role'], p['signature'], p['name']]
            table_data.append(row)
        sig_table = Table(table_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
        sig_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('FONTNAME', (0, 0), (-1, -1), 'Times New Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), people[0]['font_size']),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black)
        ]))
        elements.append(sig_table)
        elements.append(Spacer(1, 12))

    # Appendices
    for appendix in data['appendices']:
        elements.append(PageBreak())
        text = format_text(appendix['title']['text'])
        elements.append(Paragraph(text, get_paragraph_style(appendix['title']['font_size'], "center")))
        elements.append(Spacer(1, 12))
        if isinstance(appendix['content'], dict) and 'image' in appendix['content']:
            add_centered_image(elements, appendix['content']['image']['path'])
        else:
            text = format_text(appendix['content']['text'])
            elements.append(Paragraph(text, get_paragraph_style(appendix['content']['font_size'])))

    doc.build(elements)

def create_docx_report(data, output_path):
    doc = Document()
    doc.styles.add_style('Custom List Number', WD_STYLE_TYPE.LIST)

    def add_formatted_paragraph(doc, text, font_size, alignment="left"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if alignment == "center" else WD_ALIGN_PARAGRAPH.LEFT
        parts = re.split(r'(<b>.*?</b>)', text)
        for part in parts:
            if part.startswith('<b>') and part.endswith('</b>'):
                run = p.add_run(part[3:-4])
                run.bold = True
            else:
                run = p.add_run(part)
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'
        return p

    def add_formatted_cell(cell, text, font_size):
        cell.text = ""
        parts = re.split(r'(<b>.*?</b>)', text)
        for part in parts:
            if part.startswith('<b>') and part.endswith('</b>'):
                run = cell.paragraphs[0].add_run(part[3:-4])
                run.bold = True
            else:
                run = cell.paragraphs[0].add_run(part)
            run.font.size = Pt(font_size)
            run.font.name = 'Times New Roman'

    def add_centered_picture(doc, img_path):
        if os.path.exists(img_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path, width=Inches(450 / 72))
        else:
            add_formatted_paragraph(doc, f"[Изображение не найдено: {img_path}]", 10)

    # Title and subtitle
    for item in [data['title'], data['subtitle']]:
        add_formatted_paragraph(doc, item['text'], item['font_size'], item['alignment'])
    doc.add_paragraph("")

    # Sections
    for section in data['sections']:
        add_formatted_paragraph(doc, f"{section['id']}. <b>{section['title']['text']}</b>", section['title']['font_size'])
        if 'content' in section:
            if isinstance(section['content'], dict) and 'text' in section['content']:
                add_formatted_paragraph(doc, section['content']['text'], section['content']['font_size'])
            elif isinstance(section['content'], list):
                for item in section['content']:
                    if isinstance(item, dict) and 'image' in item:
                        add_centered_picture(doc, item['image']['path'])
                        add_formatted_paragraph(doc, item['text']['text'], item['text']['font_size'], "center")
                    else:
                        add_formatted_paragraph(doc, item['text'], item['font_size'])

        if 'table' in section:
            table = doc.add_table(rows=1, cols=len(section['table']['headers']))
            table.style = 'Table Grid'
            hdr_row = table.rows[0]
            for i, header in enumerate(section['table']['headers']):
                add_formatted_cell(hdr_row.cells[i], header, section['table']['font_size'])
            for row_data in section['table']['rows']:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row_data):
                    add_formatted_cell(row_cells[i], str(cell), section['table']['font_size'])

        if 'subsections' in section:
            for subsection in section['subsections']:
                add_formatted_paragraph(doc, subsection['title']['text'], subsection['title']['font_size'])
                if 'table' in subsection:
                    table = doc.add_table(rows=1, cols=len(subsection['table']['headers']))
                    table.style = 'Table Grid'
                    hdr_row = table.rows[0]
                    for i, header in enumerate(subsection['table']['headers']):
                        add_formatted_cell(hdr_row.cells[i], header, subsection['table']['font_size'])
                    for row_data in subsection['table']['rows']:
                        row_cells = table.add_row().cells
                        for i, cell in enumerate(row_data):
                            add_formatted_cell(row_cells[i], str(cell), subsection['table']['font_size'])
                elif 'content' in subsection:
                    add_formatted_paragraph(doc, subsection['content']['text'], subsection['content']['font_size'])

    # Developers and Approvers Table
    for title_key, people in [("РАЗРАБОТАНО:", data['developers']), ("СОГЛАСОВАНО:", data['approvers'])]:
        add_formatted_paragraph(doc, title_key, 10)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        headers = ["<b>Должность</b>", "<b>Подпись</b>", "<b>Ф.И.О.</b>"]
        hdr_row = table.rows[0]
        for i, header in enumerate(headers):
            add_formatted_cell(hdr_row.cells[i], header, people[0]['font_size'])
        for person in people:
            row = table.add_row().cells
            add_formatted_cell(row[0], person['role'], person['font_size'])
            add_formatted_cell(row[1], person['signature'], person['font_size'])
            add_formatted_cell(row[2], person['name'], person['font_size'])

    # Appendices
    for appendix in data['appendices']:
        doc.add_page_break()
        add_formatted_paragraph(doc, appendix['title']['text'], appendix['title']['font_size'], "center")
        if isinstance(appendix['content'], dict) and 'image' in appendix['content']:
            add_centered_picture(doc, appendix['content']['image']['path'])
        else:
            add_formatted_paragraph(doc, appendix['content']['text'], appendix['content']['font_size'])

    doc.save(output_path)

def shorten_date(date_str):
    """Shorten date from dd.mm.yyyy to dd.mm.yy"""
    if len(date_str) == 10:
        return date_str[:6] + date_str[8:]
    return date_str

def generate_report(data, pdf_output=None, docx_output=None):
    start_short = shorten_date(data['period']['start_date'])
    end_short = shorten_date(data['period']['end_date'])
    base_name = f"Отчет_{data['report_type']}_{start_short}-{end_short}"
    
    pdf_output = pdf_output or f"{base_name}.pdf"
    docx_output = docx_output or f"{base_name}.docx"
    
    create_pdf_report(data, pdf_output)
    create_docx_report(data, docx_output)
    return pdf_output, docx_output

if __name__ == "__main__":
    pdf_path, docx_path = generate_report(build_report_data(sample_minimal_json))
    print(f"PDF: {pdf_path}, DOCX: {docx_path}")
